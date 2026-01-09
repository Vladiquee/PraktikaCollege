import tkinter as tk
from tkinter import messagebox, ttk
import customtkinter as ctk
import sqlite3
import random
import string
import re
import os
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import pandas as pd

# Налаштування вигляду
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class InternshipSystem(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("АРМ: Система обліку практики ДНЗ «Полтавський політехнічний ліцей»")
        self.geometry("1350x900")
        self.minsize(1050, 750)
        
        self.current_user = None 
        self.init_db() 
        
        self.container = ctk.CTkFrame(self, fg_color="transparent")
        self.container.pack(side="top", fill="both", expand=True)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)

        self.frames = {}
        for F in (LoginFrame, RegisterFrame, RecoveryFrame, MainAppFrame):
            page_name = F.__name__
            frame = F(parent=self.container, controller=self)
            self.frames[page_name] = frame
            frame.grid(row=0, column=0, sticky="nsew")
        self.show_frame("LoginFrame")

    def init_db(self):
        # Використовуємо назву бази даних згідно з техзавданням
        conn = sqlite3.connect('internship_system.db')
        cursor = conn.cursor()
        cursor.execute("PRAGMA foreign_keys = ON;")
        cursor.execute('''CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, surname TEXT, login TEXT UNIQUE, password TEXT, role TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS students (id INTEGER PRIMARY KEY AUTOINCREMENT, fullname TEXT, specialty TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS companies (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, address TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS internships (id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER, company_id INTEGER, start_date TEXT, end_date TEXT, status TEXT, grade INTEGER, FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE, FOREIGN KEY(company_id) REFERENCES companies(id) ON DELETE CASCADE)''')
        conn.commit()
        conn.close()

    def show_frame(self, page_name):
        frame = self.frames[page_name]
        if page_name == "MainAppFrame": frame.build_menu()
        frame.tkraise()

# --- ЕКРАН ВХОДУ ---
class LoginFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        card = ctk.CTkFrame(self, width=400, height=550, corner_radius=20, border_width=2)
        card.place(relx=0.5, rely=0.5, anchor="center")
        
        ctk.CTkLabel(card, text="🔑", font=("Arial", 60)).pack(pady=(40, 10))
        ctk.CTkLabel(card, text="АВТОРИЗАЦІЯ", font=("Arial", 24, "bold")).pack(pady=10)
        
        self.l_e = ctk.CTkEntry(card, width=300, height=45, placeholder_text="Логін")
        self.l_e.pack(pady=10)
        
        pass_f = ctk.CTkFrame(card, fg_color="transparent")
        pass_f.pack(pady=10)
        self.p_e = ctk.CTkEntry(pass_f, width=255, height=45, placeholder_text="Пароль", show="*")
        self.p_e.pack(side="left")
        self.eye = ctk.CTkButton(pass_f, text="👁", width=40, height=45, fg_color="#333", command=self.toggle)
        self.eye.pack(side="left", padx=5)
        
        ctk.CTkButton(card, text="УВІЙТИ", width=300, height=50, command=self.login).pack(pady=20)
        ctk.CTkButton(card, text="Створити аккаунт", fg_color="transparent", command=lambda: controller.show_frame("RegisterFrame")).pack()
        ctk.CTkButton(card, text="Забули дані?", fg_color="transparent", text_color="gray", command=lambda: controller.show_frame("RecoveryFrame")).pack()

    def toggle(self):
        s = "" if self.p_e.cget("show") == "*" else "*"
        self.p_e.configure(show=s); self.eye.configure(text="🔒" if s == "" else "👁")

    def login(self):
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
        cur.execute("SELECT name, surname, login, role FROM users WHERE login=? AND password=?", (self.l_e.get(), self.p_e.get()))
        u = cur.fetchone(); conn.close()
        if u:
            self.controller.current_user = {"name": u[0], "surname": u[1], "login": u[2], "role": u[3]}
            self.controller.show_frame("MainAppFrame")
        else: messagebox.showerror("Помилка", "Невірні дані")

# --- ЕКРАН РЕЄСТРАЦІЇ ---
class RegisterFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        card = ctk.CTkFrame(self, width=450, height=650, corner_radius=20, border_width=2)
        card.place(relx=0.5, rely=0.5, anchor="center")
        ctk.CTkLabel(card, text="📝 РЕЄСТРАЦІЯ", font=("Arial", 24, "bold")).pack(pady=20)
        
        self.n = ctk.CTkEntry(card, width=350, placeholder_text="Ім'я (Latin)"); self.n.pack(pady=5)
        self.s = ctk.CTkEntry(card, width=350, placeholder_text="Прізвище (Latin)"); self.s.pack(pady=5)
        
        pf = ctk.CTkFrame(card, fg_color="transparent"); pf.pack(pady=5)
        self.p = ctk.CTkEntry(pf, width=305, placeholder_text="Пароль (6+)", show="*"); self.p.pack(side="left")
        self.eye = ctk.CTkButton(pf, text="👁", width=40, fg_color="#333", command=self.toggle); self.eye.pack(side="left", padx=5)

        self.role_v = ctk.StringVar(value="Учень")
        ctk.CTkComboBox(card, values=["Адмін", "Вчитель", "Учень"], variable=self.role_v, width=350).pack(pady=10)
        self.rp_e = ctk.CTkEntry(card, width=350, placeholder_text="Код доступу", show="*"); self.rp_e.pack(pady=5)
        
        ctk.CTkButton(card, text="ЗАРЕЄСТРУВАТИСЯ", width=350, height=45, command=self.reg).pack(pady=20)
        ctk.CTkButton(card, text="Назад", fg_color="transparent", command=lambda: controller.show_frame("LoginFrame")).pack()

    def toggle(self):
        s = "" if self.p.cget("show") == "*" else "*"
        self.p.configure(show=s); self.eye.configure(text="🔒" if s == "" else "👁")

    def reg(self):
        n, s, p, r, rp = self.n.get().strip(), self.s.get().strip(), self.p.get().strip(), self.role_v.get(), self.rp_e.get()
        if not re.match(r"^[A-Za-z]+$", n) or not re.match(r"^[A-Za-z]+$", s): return messagebox.showwarning("Помилка", "Тільки латиниця!")
        if len(p) < 6: return messagebox.showwarning("Помилка", "Короткий пароль!")
        if r == "Адмін" and rp != "123456789": return messagebox.showerror("Доступ", "Невірний код Адміна!")
        if r == "Вчитель" and rp != "987654321": return messagebox.showerror("Доступ", "Невірний код Вчителя!")
        
        login = f"{s[:3].upper()}{''.join(random.choices(string.digits, k=4))}"
        try:
            conn = sqlite3.connect('internship_system.db')
            conn.execute("INSERT INTO users (name, surname, login, password, role) VALUES (?,?,?,?,?)", (n, s, login, p, r))
            conn.commit(); conn.close()
            messagebox.showinfo("Успіх", f"Аккаунт створено! Ваш логін: {login}"); self.controller.show_frame("LoginFrame")
        except: messagebox.showerror("Помилка", "Логін вже зайнятий")

# --- ЕКРАН ВІДНОВЛЕННЯ ---
class RecoveryFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        card = ctk.CTkFrame(self, width=450, height=600, corner_radius=20, border_width=2)
        card.place(relx=0.5, rely=0.5, anchor="center")
        ctk.CTkLabel(card, text="🔄 ВІДНОВЛЕННЯ", font=("Arial", 24, "bold")).pack(pady=20)
        self.e1 = ctk.CTkEntry(card, width=350, placeholder_text="Ваше ім'я"); self.e1.pack(pady=5)
        self.e2 = ctk.CTkEntry(card, width=350, placeholder_text="Ваше прізвище"); self.e2.pack(pady=5)
        ctk.CTkButton(card, text="Знайти логін", command=self.get_l).pack(pady=10)
        ctk.CTkLabel(card, text="⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯", text_color="#333").pack(pady=10)
        self.e3 = ctk.CTkEntry(card, width=350, placeholder_text="Ваш логін"); self.e3.pack(pady=5)
        self.e4 = ctk.CTkEntry(card, width=350, placeholder_text="Ваше прізвище"); self.e4.pack(pady=5)
        ctk.CTkButton(card, text="Відновити пароль", command=self.get_p).pack(pady=10)
        ctk.CTkButton(card, text="Назад", fg_color="transparent", command=lambda: controller.show_frame("LoginFrame")).pack()

    def get_l(self):
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
        cur.execute("SELECT login FROM users WHERE name=? AND surname=?", (self.e1.get(), self.e2.get())); r = cur.fetchone(); conn.close()
        if r: messagebox.showinfo("Результат", f"Ваш логін: {r[0]}")
        else: messagebox.showerror("Помилка", "Не знайдено")
    def get_p(self):
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
        cur.execute("SELECT password FROM users WHERE login=? AND surname=?", (self.e3.get(), self.e4.get())); r = cur.fetchone(); conn.close()
        if r: messagebox.showinfo("Результат", f"Ваш пароль: {r[0]}")
        else: messagebox.showerror("Помилка", "Дані не збігаються")

# --- ГОЛОВНА РОБОЧА ПАНЕЛЬ ---
class MainAppFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        self.sidebar = ctk.CTkFrame(self, width=260, corner_radius=0, border_width=1); self.sidebar.pack(side="left", fill="y")
        self.content = ctk.CTkScrollableFrame(self, fg_color="transparent", corner_radius=0); self.content.pack(side="right", expand=True, fill="both", padx=10, pady=10)

    def build_menu(self):
        for w in self.sidebar.winfo_children(): w.destroy()
        u = self.controller.current_user
        ctk.CTkLabel(self.sidebar, text=f"👤 {u['role'].upper()}", font=("Arial", 18, "bold"), text_color="cyan").pack(pady=30)
        
        btns = [("Мій кабінет", self.show_profile), ("Учні ліцею", self.show_students), ("Підприємства", self.show_companies), 
                ("Практика", self.show_internships), ("Аналітика", self.show_stats), ("Допомога", self.show_help)]
        
        for t, c in btns:
            if u['role'] == "Учень" and t != "Мій кабінет": continue
            ctk.CTkButton(self.sidebar, text=t, height=45, command=c).pack(pady=5, padx=15, fill="x")
        
        if u['role'] == "Учень": ctk.CTkButton(self.sidebar, text="📊 Моя оцінка", fg_color="green", command=self.show_grade).pack(pady=5, padx=15, fill="x")
        ctk.CTkButton(self.sidebar, text="🚪 Вийти", fg_color="red", command=lambda: self.controller.show_frame("LoginFrame")).pack(side="bottom", pady=20, padx=15, fill="x")
        self.show_profile()

    def clear(self):
        for w in self.content.winfo_children(): w.destroy()

    def show_profile(self):
        self.clear(); u = self.controller.current_user
        ctk.CTkLabel(self.content, text="ОСОБИСТИЙ КАБІНЕТ", font=("Arial", 28, "bold")).pack(pady=20)
        box = ctk.CTkFrame(self.content, corner_radius=15, border_width=1); box.pack(pady=10, fill="x", padx=20)
        ctk.CTkLabel(box, text=f"Користувач: {u['name']} {u['surname']}", font=("Arial", 18)).pack(pady=10)
        ctk.CTkLabel(box, text=f"ЛОГІН: {u['login']}", font=("Arial", 20, "bold"), text_color="yellow").pack(pady=10)
        
        if u['role'] == "Адмін":
            ctk.CTkButton(self.content, text="💾 Резервна копія БД", fg_color="#34495e", command=self.backup).pack(pady=20)

    def backup(self):
        if not os.path.exists('backups'): os.makedirs('backups')
        p = f"backups/db_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
        shutil.copyfile('internship_system.db', p); messagebox.showinfo("Успіх", f"Копія: {p}")

    def show_students(self):
        self.clear(); u = self.controller.current_user
        ctk.CTkLabel(self.content, text="БАЗА УЧНІВ", font=("Arial", 24, "bold")).pack(pady=10)
        
        sf = ctk.CTkFrame(self.content); sf.pack(fill="x", pady=5)
        self.se = ctk.CTkEntry(sf, placeholder_text="🔍 Пошук за прізвищем...", width=400)
        self.se.grid(row=0, column=0, padx=10, pady=10); self.se.bind("<KeyRelease>", self.filter_st)

        f = ctk.CTkFrame(self.content); f.pack(fill="x", pady=10)
        en, es = ctk.CTkEntry(f, placeholder_text="ПІБ"), ctk.CTkEntry(f, placeholder_text="Спеціальність")
        en.grid(row=0, column=0, padx=5, pady=10); es.grid(row=0, column=1, padx=5)
        def add():
            conn = sqlite3.connect('internship_system.db'); conn.execute("INSERT INTO students (fullname, specialty) VALUES (?,?)", (en.get(), es.get())); conn.commit(); self.show_students()
        ctk.CTkButton(f, text="Додати", width=100, command=add).grid(row=0, column=2, padx=5)
        if u['role'] == "Адмін":
            ctk.CTkButton(f, text="Видалити", fg_color="red", width=100, command=lambda: self.del_r("students", self.tr_st, self.show_students)).grid(row=0, column=3, padx=5)
        ctk.CTkButton(self.content, text="📊 Експорт Excel", fg_color="green", command=self.export_ex).pack(pady=5)
        self.tr_st = self.render_table(("ID", "ПІБ", "Спеціальність"), "SELECT * FROM students")

    def filter_st(self, e):
        t = self.se.get()
        for i in self.tr_st.get_children(): self.tr_st.delete(i)
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
        cur.execute("SELECT * FROM students WHERE fullname LIKE ?", ('%'+t+'%',))
        for row in cur.fetchall(): self.tr_st.insert("", "end", values=row)
        conn.close()

    def show_companies(self):
        self.clear(); u = self.controller.current_user
        ctk.CTkLabel(self.content, text="ПІДПРИЄМСТВА", font=("Arial", 24, "bold")).pack(pady=10)
        f = ctk.CTkFrame(self.content); f.pack(fill="x", pady=10)
        en, ea = ctk.CTkEntry(f, placeholder_text="Назва"), ctk.CTkEntry(f, placeholder_text="Адреса")
        en.grid(row=0, column=0, padx=5, pady=10); ea.grid(row=0, column=1, padx=5)
        def add():
            conn = sqlite3.connect('internship_system.db'); conn.execute("INSERT INTO companies (name, address) VALUES (?,?)", (en.get(), ea.get())); conn.commit(); self.show_companies()
        ctk.CTkButton(f, text="Зберегти", width=100, command=add).grid(row=0, column=2, padx=5)
        if u['role'] == "Адмін":
            ctk.CTkButton(f, text="Видалити", fg_color="red", width=100, command=lambda: self.del_r("companies", self.tr_co, self.show_companies)).grid(row=0, column=3, padx=5)
        self.tr_co = self.render_table(("ID", "Назва", "Адреса"), "SELECT * FROM companies")

    def show_internships(self):
        self.clear(); u = self.controller.current_user
        ctk.CTkLabel(self.content, text="ЖУРНАЛ ПРАКТИКИ", font=("Arial", 24, "bold")).pack(pady=10)
        f = ctk.CTkFrame(self.content); f.pack(fill="x", pady=10)
        sid, cid = ctk.CTkEntry(f, width=50, placeholder_text="ID Ст"), ctk.CTkEntry(f, width=50, placeholder_text="ID Пд")
        sid.grid(row=0, column=0, padx=2, pady=10); cid.grid(row=0, column=1, padx=2)
        ed = ctk.CTkComboBox(f, values=[str(i).zfill(2) for i in range(1,32)], width=60); ed.grid(row=0, column=2)
        em = ctk.CTkComboBox(f, values=[str(i).zfill(2) for i in range(1,13)], width=60); em.grid(row=0, column=3)
        ey = ctk.CTkComboBox(f, values=["2025", "2026"], width=80); ey.grid(row=0, column=4)
        st, gr = ctk.StringVar(value="В процесі"), ctk.StringVar(value="5")
        ctk.CTkComboBox(f, values=["В процесі", "Завершено"], variable=st, width=110).grid(row=0, column=5, padx=2)
        ctk.CTkComboBox(f, values=["1","2","3","4","5"], variable=gr, width=50).grid(row=0, column=6, padx=2)
        def assign():
            fe = f"{ey.get()}-{em.get()}-{ed.get()}"
            conn = sqlite3.connect('internship_system.db'); conn.execute("INSERT INTO internships (student_id, company_id, end_date, status, grade) VALUES (?,?,?,?,?)", (sid.get(), cid.get(), fe, st.get(), gr.get())); conn.commit(); self.show_internships()
        ctk.CTkButton(f, text="Оцінити", width=80, command=assign).grid(row=0, column=7, padx=5)
        ctk.CTkButton(self.content, text="📄 Word Звіт", width=120, command=self.rep_docx).pack(pady=5)
        self.tr_in = self.render_table(("№", "Студент", "Компанія", "Кінець", "Статус", "Оцінка"), "SELECT i.id, s.fullname, c.name, i.end_date, i.status, i.grade FROM internships i JOIN students s ON i.student_id = s.id JOIN companies c ON i.company_id = c.id")

    def show_stats(self):
        self.clear()
        conn = sqlite3.connect('internship_system.db')
        df1 = pd.read_sql_query("SELECT grade, COUNT(*) as count FROM internships GROUP BY grade", conn)
        df2 = pd.read_sql_query("SELECT s.specialty, AVG(i.grade) as avg FROM internships i JOIN students s ON i.student_id = s.id GROUP BY s.specialty", conn); conn.close()
        if not df1.empty:
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(8, 10)); fig.patch.set_facecolor('#2b2b2b')
            ax1.set_facecolor('#2b2b2b'); ax1.bar(df1['grade'].astype(str), df1['count'], color='#1f538d'); ax1.tick_params(colors='white')
            ax2.set_facecolor('#2b2b2b'); ax2.barh(df2['specialty'], df2['avg'], color='#2ecc71'); ax2.tick_params(colors='white')
            plt.tight_layout(); canvas = FigureCanvasTkAgg(fig, master=self.content); canvas.draw(); canvas.get_tk_widget().pack(fill="x")

    def show_help(self):
        self.clear()
        ctk.CTkLabel(self.content, text="ДОВІДКА АРМ", font=("Arial", 24, "bold")).pack(pady=20)
        t = ("- Реєстрація Адміна: код 123456789, Вчителя: 987654321.\n"
             "- Видалення: Доступне лише Адміну. Перераховує ID автоматично.\n"
             "- Аналітика: Розраховує середній бал за спеціальностями.\n"
             "- Бекап: Доступний в кабінеті Адміна.")
        ctk.CTkLabel(self.content, text=t, justify="left", font=("Arial", 14)).pack(padx=20)

    def del_r(self, tbl, tr, ref):
        sel = tr.selection()
        if sel:
            rid = tr.item(sel[0])['values'][0]
            conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
            cur.execute(f"DELETE FROM {tbl} WHERE id=?"); cur.execute(f"UPDATE {tbl} SET id = id - 1 WHERE id > ?", (rid,))
            cur.execute(f"UPDATE sqlite_sequence SET seq = (SELECT COUNT(*) FROM {tbl}) WHERE name = '{tbl}'")
            conn.commit(); conn.close(); ref()

    def rep_docx(self):
        # Оновлений метод генерації звіту у "Завантаження"
        sel = self.tr_in.selection()
        if not sel:
            messagebox.showwarning("Вибір", "Будь ласка, виберіть студента з таблиці!")
            return

        try:
            r = self.tr_in.item(sel[0])['values']
            student_name, company_name, grade = str(r[1]), str(r[2]), str(r[5])
            downloads_path = str(Path.home() / "Downloads")
            safe_name = "".join([c for c in student_name if c.isalnum() or c in (' ', '_')]).rstrip()
            file_name = f"Zvit_{safe_name}.docx"
            full_path = os.path.join(downloads_path, file_name)
            
            # Шукаємо шаблон поруч із програмою
            template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx")
            
            if os.path.exists(template_path):
                doc = Document(template_path)
                for p in doc.paragraphs:
                    if "{student}" in p.text: p.text = p.text.replace("{student}", student_name)
                    if "{company}" in p.text: p.text = p.text.replace("{company}", company_name)
                    if "{grade}" in p.text: p.text = p.text.replace("{grade}", grade)
            else:
                doc = Document()
                doc.add_heading('Звіт про виробничу практику', 0)
                doc.add_paragraph(f"Студент: {student_name}")
                doc.add_paragraph(f"Компанія: {company_name}")
                doc.add_paragraph(f"Оцінка: {grade}")
            
            doc.save(full_path)
            messagebox.showinfo("Успіх", f"Звіт збережено у Завантаження:\n{file_name}")
            os.startfile(downloads_path) # Відкриваємо папку для зручності
        except Exception as e:
            messagebox.showerror("Помилка", f"Не вдалося створити звіт: {str(e)}")

    def export_ex(self):
        conn = sqlite3.connect('internship_system.db'); df = pd.read_sql_query("SELECT * FROM students", conn); conn.close()
        df.to_excel("Студенти_ППЛ.xlsx", index=False); messagebox.showinfo("Excel", "Експортовано успішно")

    def show_grade(self):
        self.clear(); u = self.controller.current_user
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor()
        cur.execute("SELECT c.name, i.status, i.grade FROM internships i JOIN students s ON i.student_id = s.id JOIN companies c ON i.company_id = c.id WHERE s.fullname LIKE ?", ('%'+u['surname']+'%',))
        r = cur.fetchone(); conn.close()
        if r:
            box = ctk.CTkFrame(self.content, border_width=2, border_color="green"); box.pack(pady=20, padx=20)
            ctk.CTkLabel(box, text=f"📍 {r[0]}", font=("Arial", 18)).pack(pady=10); ctk.CTkLabel(box, text=f"ОЦІНКА: {r[2]}", font=("Arial", 45, "bold"), text_color="green").pack(pady=20)
        else: ctk.CTkLabel(self.content, text="Результатів ще немає.").pack()

    def render_table(self, h, q):
        fr = ctk.CTkFrame(self.content); fr.pack(expand=True, fill="both", pady=10); tree = ttk.Treeview(fr, columns=h, show="headings")
        style = ttk.Style(); style.theme_use("default"); style.configure("Treeview", background="#2b2b2b", foreground="white", fieldbackground="#2b2b2b", rowheight=35); style.map("Treeview", background=[("selected", "#1f538d")])
        for x in h: tree.heading(x, text=x); tree.column(x, width=130, anchor="center")
        conn = sqlite3.connect('internship_system.db'); cur = conn.cursor(); cur.execute(q)
        for row in cur.fetchall(): tree.insert("", "end", values=row)
        conn.close(); tree.pack(expand=True, fill="both"); return tree

if __name__ == "__main__":
    app = InternshipSystem(); app.mainloop()
